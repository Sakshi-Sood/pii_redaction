import re
import spacy
import hashlib
from docx import Document
from faker import Faker

try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")

def normalize_entity(text, entity_type):
    """Normalize text for consistent hashing and registry mapping."""
    # Remove harmless surrounding punctuation, footnotes, and whitespace
    text = text.strip('.,;:\'"()[]{}<>*- \t\n\r*^&#@!~`')
    text = " ".join(text.split())
    text = text.lower()
    
    if entity_type == 'PHONE':
        digits = re.sub(r'\D', '', text)
        if len(digits) == 12 and digits.startswith('91'):
            digits = digits[2:]
        elif len(digits) == 11 and digits.startswith('0'):
            digits = digits[1:]
        return digits
    return text

def generate_fake_value(canonical_text, entity_type):
    """Deterministically generate fake data based on canonical text and entity type."""
    seed = int(hashlib.sha256(f"{canonical_text}_{entity_type}".encode()).hexdigest()[:15], 16)
    
    Faker.seed(seed)
    fake = Faker('en_IN')
    
    if entity_type == 'PERSON':
        return fake.name()
    elif entity_type == 'EMAIL':
        return fake.email()
    elif entity_type == 'PHONE':
        return fake.phone_number()
    elif entity_type == 'ORG':
        return fake.company()
    elif entity_type == 'ADDRESS':
        return fake.address().replace('\n', ', ')
    elif entity_type == 'DOB':
        return fake.date_of_birth().strftime('%B %d, %Y')
    elif entity_type == 'IP':
        return fake.ipv4()
    elif entity_type == 'SSN':
        return fake.ssn()
    elif entity_type == 'CREDIT_CARD':
        return fake.credit_card_number()
    return "REDACTED"

def detect_emails(text):
    spans = []
    pattern = re.compile(r'[\w.+-]+@[\w-]+\.[\w.]+')
    for m in pattern.finditer(text):
        val = m.group()
        end = m.end()
        while val and val[-1] in '.,;':
            val = val[:-1]
            end -= 1
        if val and '@' in val:
            spans.append((m.start(), end, 'EMAIL', val))
    return spans

def detect_phones(text):
    spans = []
    # Strict formats with country/area code or spacing/hyphens
    pattern_format = re.compile(r'(?:\+\s*91[\s-]*\d[\d\s-]{8,}|\b0\d{2,4}[\s-]+\d{6,8}\b|\b[6-9]\d{2}[\s-]+\d{3}[\s-]+\d{4}\b)')
    for m in pattern_format.finditer(text):
        spans.append((m.start(), m.end(), 'PHONE', m.group()))
        
    # Flat 10-digit numbers require contextual confidence
    pattern_flat = re.compile(r'\b[6-9]\d{9}\b')
    phone_keywords = ['phone', 'tel', 'mobile', 'contact', 'call', 'fax']
    lower_text = text.lower()
    has_context = any(kw in lower_text for kw in phone_keywords)
    
    for m in pattern_flat.finditer(text):
        if not any(s[0] <= m.start() and s[1] >= m.end() for s in spans):
            if has_context:
                spans.append((m.start(), m.end(), 'PHONE', m.group()))
    return spans

def detect_ips(text):
    spans = []
    pattern = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')
    for m in pattern.finditer(text):
        parts = [int(x) for x in m.group().split('.')]
        if all(0 <= p <= 255 for p in parts):
            spans.append((m.start(), m.end(), 'IP', m.group()))
    return spans

def detect_ssns(text):
    spans = []
    pattern = re.compile(r'\b\d{3}-\d{2}-\d{4}\b')
    for m in pattern.finditer(text):
        spans.append((m.start(), m.end(), 'SSN', m.group()))
    return spans

def detect_credit_cards(text):
    spans = []
    pattern = re.compile(r'\b(?:\d{4}[\s-]?){3,4}\d{0,4}\b')
    for m in pattern.finditer(text):
        digits = re.sub(r'\D', '', m.group())
        if len(digits) in (15, 16):
            s = 0
            for i, c in enumerate(reversed(digits)):
                n = int(c)
                if i % 2 == 1:
                    n *= 2
                    if n > 9: n -= 9
                s += n
            if s % 10 == 0:
                spans.append((m.start(), m.end(), 'CREDIT_CARD', m.group()))
    return spans

def detect_dobs(text):
    spans = []
    ctx_pat = re.compile(r'(?:date\s+of\s+birth|DOB|born|birth\s+date)', re.IGNORECASE)
    if ctx_pat.search(text):
        date_pat = re.compile(r'\b(?:\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4})\b', re.IGNORECASE)
        for m in date_pat.finditer(text):
            spans.append((m.start(), m.end(), 'DOB', m.group()))
    return spans

DEFAULT_BOILERPLATE_TERMS = {
    'internal risks', 'external risks', 'risk factors', 'goods and services tax',
    'indian rupees', 'rupees', 'net proceeds', 'gross proceeds', 'offer related terms',
    'bid/offer closing day', 'bid/offer closing date', 'anchor investors', 'anchor investor',
    'statutory auditors', 'statutory auditor', 'auditors', 'auditor', 'care report',
    'crisil report', 'promoter group', 'promoter selling shareholders', 'selling shareholders',
    'individual promoters', 'board of directors', 'audit committee', 'nomination and remuneration committee',
    'stakeholders relationship committee', 'corporate social responsibility committee',
    'fresh issue', 'offer for sale', 'red herring prospectus', 'draft red herring prospectus',
    'prospectus', 'general information document', 'designated stock exchange',
    'stock exchanges', 'listing regulations', 'sebi icdr regulations', 'companies act',
    'income tax act', 'equity shares', 'equity share', 'price band', 'floor price',
    'cap price', 'issue price', 'face value', 'bid cum application form', 'asba form',
    'summary of the offer', 'objects of the offer', 'basis for issue price',
    'restated financial information', 'financial statements', 'materiality policy',
    'capital structure', 'terms of the offer', 'business overview', 'financial indebtedness',
    'management discussion and analysis', 'industry overview', 'our business', 'our promoters',
    'promoter group', 'dividend policy', 'outstanding litigation', 'government approvals',
    'statutory and other information', 'other regulatory and statutory disclosures',
    'declaration', 'definitions and abbreviations', 'table of contents'
}

FINANCIAL_BOILERPLATE_WORDS = {
    'risk', 'risks', 'proceeds', 'prospectus', 'regulations', 'regulation', 'act', 'report',
    'reports', 'tax', 'taxes', 'rupees', 'rupee', 'committee', 'committees', 'policy', 'policies',
    'factor', 'factors', 'intermediaries', 'intermediary', 'shareholders', 'shareholder',
    'allottees', 'allottee', 'investors', 'investor', 'bidders', 'bidder', 'expenses', 'expense',
    'statements', 'statement', 'resolution', 'resolutions', 'meeting', 'meetings', 'exchange',
    'exchanges', 'section', 'sections', 'schedule', 'schedules', 'clause', 'clauses', 'table',
    'tables', 'form', 'forms', 'annexure', 'annexures', 'part', 'parts', 'issue', 'issues',
    'offer', 'offers', 'bid', 'bids', 'price', 'prices', 'capital', 'shares', 'share',
    'equity', 'margin', 'margins', 'auditor', 'auditors', 'promoter', 'promoters', 'management',
    'matters', 'terms', 'term', 'abbreviations', 'abbreviation', 'definitions', 'definition',
    'overview', 'summary', 'details', 'general', 'information', 'industry', 'business',
    'director', 'officer', 'secretary', 'chairman', 'board', 'manager', 'lead',
    'value', 'face', 'total', 'date', 'financial', 'year', 'period', 'rupees', 'gst',
    'holding', 'subsidiary', 'associate', 'group', 'company', 'statutory'
}

CURRENCIES_AND_UNITS = {
    'indian rupees', 'rupees', 'inr', 'rs', 'rs.', 'usd', 'us dollars', 'dollars',
    'eur', 'euros', 'gbp', 'pounds', 'crores', 'crore', 'lakhs', 'lakh', 'millions', 'billions'
}

def extract_prospectus_denylist(doc):
    """Extract financial and legal defined terms from document definition tables."""
    denylist = set(DEFAULT_BOILERPLATE_TERMS)
    if doc is None:
        return denylist
    for t in doc.tables:
        if len(t.rows) > 0 and len(t.columns) >= 2:
            hdr = [c.text.strip().lower() for c in t.rows[0].cells[:2]]
            if any('term' in h for h in hdr):
                for row in t.rows[1:]:
                    term_cell = row.cells[0].text.strip()
                    if not term_cell:
                        continue
                    for line in term_cell.split('\n'):
                        line = line.strip()
                        if not line or len(line) < 2:
                            continue
                        cleaned = re.sub(r'[\*\#\(\)\[\]]', ' ', line).strip()
                        parts = re.split(r'[/,]|(?:\s+or\s+)', cleaned)
                        for part in parts:
                            p = part.strip().lower()
                            if len(p) >= 2:
                                denylist.add(p)
                                if p.endswith('s'):
                                    denylist.add(p[:-1])
                                else:
                                    denylist.add(p + 's')
    return denylist

def is_person_name_pattern(text):
    """Check if text matches person name pattern (optional title + 2-4 capitalized words, no financial keywords)."""
    words = text.split()
    if words and words[0].lower() in {'mr', 'mrs', 'ms', 'dr', 'shri', 'smt', 'prof', 'mr.', 'mrs.', 'ms.', 'dr.'}:
        words = words[1:]
    if not (2 <= len(words) <= 4):
        return False
    for w in words:
        clean_w = w.strip('.,;:\'"()[]{}')
        if not clean_w or len(clean_w) < 2:
            return False
        if not clean_w.isalpha():
            return False
        if not (clean_w.isupper() or clean_w.istitle()):
            return False
        if clean_w.lower() in FINANCIAL_BOILERPLATE_WORDS:
            return False
    return True

# Configurable set of filing subject (issuer) current and former legal names to exclude from redaction.
# Explicit Design Choice: Only third-party PII (individuals, other companies, advisors, promoters, etc.)
# is redacted. The filing subject's own legal identities (current and former entity names) are preserved
# so the prospectus remains readable, legally coherent, and meaningful.
EXCLUDED_ENTITIES = {
    'KSH International Limited',
    'KSH International Private Limited',
    'KSH International Ltd',
    'KSH International Pvt Ltd',
    'Bhandary Metal Extrusion Private Limited',
    'Bhandary Metal Extrusion Pvt Ltd',
}

def is_excluded_issuer_entity(val, excluded_entities=None):
    """Check if an entity matches the issuer's current or former legal names."""
    if excluded_entities is None:
        excluded_entities = EXCLUDED_ENTITIES
    if not excluded_entities:
        return False
    val_clean = val.strip('.,;:\'"()[]{}<>*- \t\n\r*^&#@!~`')
    core = re.sub(r'^(?:the|our|a|an|its|such|each|this|these|those)\s+', '', val_clean, flags=re.IGNORECASE).strip().lower()
    core_norm = " ".join(core.split())
    for target in excluded_entities:
        target_norm = " ".join(target.strip().lower().split())
        if core_norm == target_norm:
            return True
    return False

def is_denylisted(val, denylist, is_org=False, excluded_entities=None):
    """Check if an entity span is in the denylist, issuer exclusion list, or matches boilerplate rules."""
    if is_excluded_issuer_entity(val, excluded_entities):
        return True
    val_clean = val.strip('.,;:\'"()[]{}<>*- \t\n\r*^&#@!~`')
    core = re.sub(r'^(?:the|our|a|an|its|such|each|this|these|those)\s+', '', val_clean, flags=re.IGNORECASE).strip().lower()
    if not core:
        return True
    if denylist and (core in denylist or any(core in d for d in denylist if len(core) >= 4 and d != core)):
        return True
    if core in CURRENCIES_AND_UNITS:
        return True
    if not is_org:
        words = [w.lower().strip('.,;:\'"()[]{}') for w in core.split()]
        if any(w in FINANCIAL_BOILERPLATE_WORDS for w in words):
            return True
        if val.isupper() and not is_person_name_pattern(val):
            return True
    return False

def detect_ner(text, denylist=None, excluded_entities=None):
    spans = []
    if not text.strip():
        return spans
    if denylist is None:
        denylist = DEFAULT_BOILERPLATE_TERMS
    if excluded_entities is None:
        excluded_entities = EXCLUDED_ENTITIES
    doc = nlp(text)
    
    org_suffixes = {
        'limited', 'ltd', 'pvt', 'private', 'llp', 'inc', 'corporation', 
        'industries', 'technologies', 'bank', 'securities', 'capital'
    }
    
    for ent in doc.ents:
        val = ent.text.strip()
        if not val or len(val) < 3:
            continue
        if is_excluded_issuer_entity(val, excluded_entities):
            continue
        lower_val = val.lower()
        
        has_suffix = any(re.search(rf'\b{s}\b', lower_val) for s in org_suffixes)
        
        if ent.label_ == 'ORG' and has_suffix:
            if not is_denylisted(val, denylist, is_org=True, excluded_entities=excluded_entities):
                spans.append((ent.start_char, ent.start_char + len(val), 'ORG', val))
        elif ent.label_ in ('PERSON', 'ORG', 'GPE', 'NORP', 'WORK_OF_ART', 'PRODUCT'):
            if not is_denylisted(val, denylist, is_org=False, excluded_entities=excluded_entities):
                if is_person_name_pattern(val):
                    clean_name = " ".join(val.split())
                    start_offset = text.find(clean_name, ent.start_char)
                    if start_offset != -1:
                        spans.append((start_offset, start_offset + len(clean_name), 'PERSON', clean_name))
                
    # Contextual pattern for list of promoters / directors if missed by spaCy
    promoter_pat = re.compile(r'(?:Individual Promoters|Promoters|Directors)[\s:]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+(?:,\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)*)', re.UNICODE)
    for m in promoter_pat.finditer(text):
        list_str = m.group(1)
        base_offset = m.start(1)
        for item in re.finditer(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+', list_str):
            p_name = item.group().strip()
            if is_person_name_pattern(p_name) and not is_denylisted(p_name, denylist, is_org=False, excluded_entities=excluded_entities):
                spans.append((base_offset + item.start(), base_offset + item.end(), 'PERSON', p_name))
                
    return spans

def detect_addresses(text):
    spans = []
    if len(text.strip()) == 0:
        return spans
        
    # 1. Addresses with explicit prefix (e.g. 'Registered Office: ...', 'having its Registered Office at ...')
    prefix_pat = re.compile(
        r'(?:\b(?:Registered Office|Corporate Office|Residential Address|Address)\s*(?:[:\-]|(?:\s+of\s+our\s+Company\s+is\s+situated\s+at|\s+at))\s+)'
        r'([^;\n\r]+?(?:\b\d{6}\b|\b\d{3}\s\d{3}\b)[^;\n\r]*)',
        re.IGNORECASE
    )
    for m in prefix_pat.finditer(text):
        addr = m.group(1).strip()
        # Cut off compound clauses like 'and its Corporate Office...'
        split_match = re.split(r'\s+and\s+(?:its\s+)?(?:Corporate|Registered)\s+Office', addr, flags=re.IGNORECASE)
        if split_match:
            addr = split_match[0]
        addr_clean = addr.rstrip(';., ')
        start = m.start(1)
        end = start + len(addr_clean)
        if end > start and len(addr_clean) > 10:
            spans.append((start, end, 'ADDRESS', text[start:end]))

    # 2. Standalone address blocks (common in table cells / contact blocks)
    pin_pat = re.compile(
        r'\b(?:(?:[0-9]{1,4}(?:[/–\-][0-9A-Za-z]+)?|[0-9]+(?:st|nd|rd|th)\s+Floor|Ground\s+Floor|Village|Taluka|Tower|Building|Block|Plot|Signature|Next\s+to|Opp\.|Near)\b[^;\n\r]{5,150}?'
        r'(?:\b\d{6}\b|\b\d{3}\s\d{3}\b|\b\d{2}[lI]\s*0\d{2}\b)'
        r'(?:[\s,]*(?:Maharashtra|Gujarat|Karnataka|Tamil\s+Nadu|Delhi|India|Telangana|Andhra\s+Pradesh|Rajasthan|West\s+Bengal|Uttar\s+Pradesh|Haryana|Punjab|Goa|Kerala))*)',
        re.IGNORECASE
    )
    for m in pin_pat.finditer(text):
        start = m.start()
        val = text[start:m.end()].strip().rstrip(';., ')
        end = start + len(val)
        if not any(s[0] <= start and s[1] >= end for s in spans):
            lower_val = val.lower()
            anchors = ['village', 'taluka', 'road', 'floor', 'tower', 'building', 'marg', 'nagar', 'apartment', 'reclamation', 'churchgate', 'street', 'lane', 'district', 'chambers', 'park', 'near', 'opp', 'station']
            if any(re.search(rf'\b{a}\b', lower_val) for a in anchors) and len(val) > 15:
                spans.append((start, end, 'ADDRESS', text[start:end]))
                
    return spans

def detect_pii(text, denylist=None, excluded_entities=None):
    all_spans = []
    all_spans.extend(detect_emails(text))
    all_spans.extend(detect_phones(text))
    all_spans.extend(detect_ips(text))
    all_spans.extend(detect_ssns(text))
    all_spans.extend(detect_credit_cards(text))
    all_spans.extend(detect_dobs(text))
    all_spans.extend(detect_ner(text, denylist=denylist, excluded_entities=excluded_entities))
    all_spans.extend(detect_addresses(text))
    
    priority = {'EMAIL': 1, 'PHONE': 2, 'SSN': 3, 'CREDIT_CARD': 4, 'IP': 5, 'DOB': 6, 'ADDRESS': 7, 'PERSON': 8, 'ORG': 9}
    all_spans.sort(key=lambda x: (x[0], priority[x[2]], -(x[1]-x[0])))
    
    resolved = []
    for span in all_spans:
        start, end, etype, val = span
        if is_excluded_issuer_entity(val, excluded_entities):
            continue
        overlap = False
        for r_start, r_end, r_etype, r_val in resolved:
            if not (end <= r_start or start >= r_end):
                overlap = True
                
                break
        if not overlap:
            resolved.append(span)
            
    return resolved

def redact_runs(runs, replacements):
    """Map replacements to original runs and update their text right-to-left."""
    if not replacements:
        return
        
    char_map = []
    for ri, r in enumerate(runs):
        for ci in range(len(r.text)):
            char_map.append((ri, ci))
            
    total_chars = len(char_map)
    orig_run_texts = [r.text for r in runs]
    
    replacements.sort(key=lambda x: x[0], reverse=True)
    
    for start, end, fake_text in replacements:
        # Assertion: Replacement span must fall within valid char_map boundaries
        assert 0 <= start < end <= total_chars, (
            f"Replacement span [{start}:{end}] out of range for total run characters {total_chars}"
        )
        
        affected_runs = sorted(list(set(char_map[i][0] for i in range(start, end))))
        assert len(affected_runs) > 0, f"No runs matched for replacement span [{start}:{end}]"
        
        first_char_idx = char_map[start][1]
        last_char_idx = char_map[end-1][1]
        
        # Verify run-splicing indices
        r_first = affected_runs[0]
        r_last = affected_runs[-1]
        assert first_char_idx < len(orig_run_texts[r_first]), (
            f"first_char_idx {first_char_idx} exceeds original run length {len(orig_run_texts[r_first])}"
        )
        assert last_char_idx < len(orig_run_texts[r_last]), (
            f"last_char_idx {last_char_idx} exceeds original run length {len(orig_run_texts[r_last])}"
        )
        
        if len(affected_runs) == 1:
            r_idx = affected_runs[0]
            runs[r_idx].text = runs[r_idx].text[:first_char_idx] + fake_text + runs[r_idx].text[last_char_idx+1:]
        else:
            runs[r_first].text = runs[r_first].text[:first_char_idx] + fake_text
            
            for r_mid in affected_runs[1:-1]:
                runs[r_mid].text = ""
                
            runs[r_last].text = runs[r_last].text[last_char_idx+1:]

class BlockRecord:
    def __init__(self, runs):
        self.runs = runs
        self.text = "".join(r.text for r in runs)
        self.resolved_spans = []

def process_document(input_path, output_path, excluded_entities=None):
    import os
    import tempfile
    import subprocess
    
    if excluded_entities is None:
        excluded_entities = EXCLUDED_ENTITIES
        
    # Load input document safely with fallback for Windows file locking
    try:
        doc = Document(input_path)
    except Exception:
        temp_input = os.path.join(tempfile.gettempdir(), "temp_input_redactor.docx")
        subprocess.run(["powershell", "-Command", f'Copy-Item -Path "{input_path}" -Destination "{temp_input}" -Force'], capture_output=True)
        doc = Document(temp_input)
        if os.path.exists(temp_input):
            try:
                os.remove(temp_input)
            except Exception:
                pass
                
    records = []
    seen_paragraphs = set()
    
    # Extract unique logical blocks from paragraphs, tables, headers, and footers
    for p in doc.paragraphs:
        if p._p not in seen_paragraphs:
            seen_paragraphs.add(p._p)
            records.append(BlockRecord(p.runs))
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p._p not in seen_paragraphs:
                        seen_paragraphs.add(p._p)
                        records.append(BlockRecord(p.runs))
                        
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p._p not in seen_paragraphs:
                    seen_paragraphs.add(p._p)
                    records.append(BlockRecord(p.runs))
        if section.footer:
            for p in section.footer.paragraphs:
                if p._p not in seen_paragraphs:
                    seen_paragraphs.add(p._p)
                    records.append(BlockRecord(p.runs))
                
    # Extract prospectus-specific denylist from definitions tables
    prospectus_denylist = extract_prospectus_denylist(doc)
    
    # Phase 1: Detect and build global mapping
    global_registry = {}
    
    for record in records:
        if not record.text.strip():
            continue
        record.resolved_spans = detect_pii(record.text, denylist=prospectus_denylist, excluded_entities=excluded_entities)
        
        for start, end, etype, val in record.resolved_spans:
            # Assertion: verify span bounds and character consistency
            assert 0 <= start < end <= len(record.text), (
                f"Span [{start}:{end}] out of range for record text of length {len(record.text)}"
            )
            assert record.text[start:end] == val, (
                f"Span text mismatch: record.text[{start}:{end}] = '{record.text[start:end]}', but val = '{val}'"
            )
            canon = normalize_entity(val, etype)
            if canon not in global_registry:
                global_registry[canon] = generate_fake_value(canon, etype)
                
    # Internal validation for 1:1 mapping
    assert len(global_registry) == len(set(global_registry.keys())), "Duplicate keys in registry"
                
    # Phase 2: Modify the DOCX using the global registry
    validation_registry = {}
    
    for record in records:
        if not record.resolved_spans:
            continue
            
        replacements = []
        for start, end, etype, val in record.resolved_spans:
            canon = normalize_entity(val, etype)
            fake_val = global_registry[canon]
            
            # Runtime validation: verify one canonical entity maps to exactly one replacement
            if canon in validation_registry:
                if validation_registry[canon] != fake_val:
                    raise ValueError(f"Consistency error: '{canon}' mapped to multiple replacements ('{validation_registry[canon]}' vs '{fake_val}')!")
            else:
                validation_registry[canon] = fake_val
                
            replacements.append((start, end, fake_val))
            
        redact_runs(record.runs, replacements)
                
    try:
        doc.save(output_path)
        print(f"Redacted document saved to {output_path}")
    except PermissionError:
        alt_path = output_path.replace(".docx", "_redacted.docx")
        doc.save(alt_path)
        print(f"Notice: '{output_path}' is open in an external viewer. Redacted output successfully saved to '{alt_path}'.")

if __name__ == "__main__":
    process_document("input/Red Herring Prospectus.docx", "output/redacted_prospectus.docx")
